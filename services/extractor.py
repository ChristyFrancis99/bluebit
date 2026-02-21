import asyncio
import io
from pathlib import Path
from typing import Tuple
import chardet
import structlog

logger = structlog.get_logger()


async def extract_text(file_bytes: bytes, filename: str) -> Tuple[str, dict]:
    """
    Extract plain text from uploaded file.
    Supports: .txt, .docx, .pdf, .md, and plain text.
    Returns (text, metadata_dict)
    """
    ext = Path(filename).suffix.lower()

    extractors = {
        ".txt": _extract_txt,
        ".md": _extract_txt,
        ".docx": _extract_docx,
        ".doc": _extract_docx,
        ".pdf": _extract_pdf,
    }

    extractor = extractors.get(ext, _extract_txt)

    try:
        text, meta = await asyncio.to_thread(extractor, file_bytes)
        meta["filename"] = filename
        meta["file_extension"] = ext
        meta["char_count"] = len(text)
        meta["word_count"] = len(text.split())
        return text.strip(), meta
    except Exception as e:
        logger.error("text_extraction.failed", filename=filename, error=str(e))
        # Fallback: try raw UTF-8
        try:
            return file_bytes.decode("utf-8", errors="replace").strip(), {
                "method": "raw_fallback"
            }
        except Exception:
            return "", {"error": str(e)}


def _extract_txt(file_bytes: bytes) -> Tuple[str, dict]:
    detected = chardet.detect(file_bytes)
    encoding = detected.get("encoding") or "utf-8"
    text = file_bytes.decode(encoding, errors="replace")
    return text, {"method": "text", "encoding": encoding}


def _extract_docx(file_bytes: bytes) -> Tuple[str, dict]:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    text = "\n\n".join(paragraphs)
    return text, {"method": "docx", "paragraph_count": len(paragraphs)}


def _extract_pdf(file_bytes: bytes) -> Tuple[str, dict]:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        text = "\n\n".join(pages)
        return text, {"method": "pdf", "page_count": len(reader.pages)}
    except Exception as e:
        return "", {"method": "pdf_failed", "error": str(e)}
